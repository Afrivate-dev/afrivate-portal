"""Branded DOCX companions for core team engagement and equity letters."""
import json
from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

PURPLE = RGBColor(0x8D, 0x40, 0x87)
INK = RGBColor(0x1F, 0x1F, 0x1F)
MUTED = RGBColor(0x5F, 0x5F, 0x5F)
SOFT = "F8F3F8"

ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "brand" / "afrivate-logo-long-purple.png"
OUT_DIR = ROOT / "hiring" / "engagement-letters"
DATA = json.loads((Path(__file__).resolve().parent / "content" / "equity-engagement-letters.json").read_text(encoding="utf-8"))


def set_run(run, *, size=11, bold=False, color=INK, all_caps=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.font.all_caps = all_caps


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def no_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def add_bottom_border(paragraph, color, size):
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(pBdr)


def para(doc, text, *, size=11, bold=False, color=INK, space_after=8, space_before=0, align="left", all_caps=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, all_caps=all_caps)
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    add_bottom_border(p, "EBDCEB", "6")
    run = p.add_run(text)
    set_run(run, size=11, bold=True, color=PURPLE, all_caps=True)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=11)
    return p


def write_letter(letter):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        r = fp.add_run("hr@afrivate.org  ·  portal.afrivate.org          AfriVate Technologies Ltd  ·  RC: 9210092")
        set_run(r, size=8, color=MUTED)

    header = doc.add_table(rows=1, cols=2)
    no_table_borders(header)
    c0, c1 = header.rows[0].cells
    c0.paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(1.7))
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, line in enumerate(["Official Document", "AfriVate Technologies Ltd", "RC: 9210092"]):
        target = c1.paragraphs[0] if i == 0 else c1.add_paragraph()
        target.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        target.paragraph_format.space_after = Pt(0)
        target.paragraph_format.space_before = Pt(0)
        set_run(target.add_run(line), size=9, color=MUTED)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(4)
    rule.paragraph_format.space_after = Pt(10)
    add_bottom_border(rule, "8D4087", "18")

    para(doc, "Core Team Engagement and Equity Letter", size=16, bold=True, align="center", space_after=2, all_caps=True)
    para(doc, letter["kicker"], size=10, color=MUTED, align="center", space_after=12)

    meta = [
        ("To", letter["fullName"]),
        ("From", DATA["from"]),
        ("Role", letter["role"]),
        ("Department", letter["department"]),
        ("Reports to", letter["reportsTo"]),
        ("Location", "Remote"),
        ("Start Date", DATA["startDate"]),
        ("Date", DATA["letterDate"]),
        ("Document Reference", DATA["documentReference"]),
    ]
    mt = doc.add_table(rows=len(meta), cols=2)
    no_table_borders(mt)
    for i, (k, v) in enumerate(meta):
        a, b = mt.rows[i].cells
        shade_cell(a, SOFT)
        shade_cell(b, SOFT)
        pa = a.paragraphs[0]
        pa.paragraph_format.space_after = Pt(2)
        pa.paragraph_format.space_before = Pt(2)
        set_run(pa.add_run(k), size=10, bold=True)
        pb = b.paragraphs[0]
        pb.paragraph_format.space_after = Pt(2)
        pb.paragraph_format.space_before = Pt(2)
        color = PURPLE if k == "Start Date" else MUTED
        set_run(pb.add_run(v), size=10, bold=(k == "Start Date"), color=color)

    doc.add_paragraph()
    para(doc, f"Written instrument: {DATA['documentReference']}.", size=10, color=MUTED, space_after=12)
    para(doc, f"Dear {letter['firstName']},", size=11, bold=True, space_after=8)
    para(doc, letter["opening"])

    heading(doc, "1. Nature of This Engagement")
    para(doc, DATA["clause1a"])
    para(doc, DATA["clause1b"])

    heading(doc, "2. Duration")
    para(doc, letter["duration"])

    heading(doc, "3. Working Structure and Agreed Capacity")
    para(doc, letter["working"])

    heading(doc, "4. Key Responsibilities")
    for item in letter["responsibilities"]:
        bullet(doc, item)
    para(doc, letter["classD"])

    heading(doc, "5. Compensation and Support")
    para(doc, DATA["clause5a"])
    para(doc, DATA["clause5b"])

    heading(doc, "6. Equity Participation")
    para(doc, DATA["clause6intro"])
    et = doc.add_table(rows=1 + len(DATA["clause6terms"]), cols=2)
    hdr = et.rows[0].cells
    shade_cell(hdr[0], SOFT)
    shade_cell(hdr[1], SOFT)
    set_run(hdr[0].paragraphs[0].add_run("Term"), size=9, bold=True, color=PURPLE, all_caps=True)
    set_run(hdr[1].paragraphs[0].add_run("Detail"), size=9, bold=True, color=PURPLE, all_caps=True)
    for i, (k, v) in enumerate(DATA["clause6terms"], start=1):
        a, b = et.rows[i].cells
        set_run(a.paragraphs[0].add_run(k), size=10, bold=True)
        set_run(b.paragraphs[0].add_run(v), size=10)
    para(doc, DATA["clause6close"], space_before=8)

    heading(doc, "7. Confidentiality and Data Protection")
    para(doc, letter["confidentiality"])

    heading(doc, "8. Intellectual Property")
    para(doc, DATA["clause8"])

    heading(doc, "9. Portfolio Use")
    para(doc, letter["portfolio"])

    heading(doc, "10. Ending the Engagement")
    bullet(doc, letter["clause10a"])
    bullet(doc, DATA["clause10b"])
    bullet(doc, DATA["clause10c"])

    heading(doc, "11. Governing Policies")
    para(doc, DATA["clause11"])

    heading(doc, "12. Not a Promise of Continued or Paid Employment")
    para(doc, DATA["clause12"])

    heading(doc, "13. Effect if This Relationship Is Later Recharacterised")
    para(doc, DATA["clause13"])

    heading(doc, "14. Acceptance")
    para(doc, DATA["clause14"])
    para(doc, "Sincerely,", space_after=18)

    sig = doc.add_table(rows=1, cols=2)
    no_table_borders(sig)
    left, right = sig.rows[0].cells
    set_run(left.paragraphs[0].add_run("Joshua Oluwasujibomi Komolafe"), size=11, bold=True)
    for line in ["Chief Executive Officer", "AfriVate Technologies Ltd"]:
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(line), size=10, color=MUTED)

    set_run(right.paragraphs[0].add_run(f"Accepted by: {letter['fullName']}"), size=11, bold=True)
    for label in ["Signature", "Date"]:
        p = right.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run("________________________________"), size=11, color=MUTED)
        p2 = right.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        set_run(p2.add_run(label), size=9, color=MUTED)

    para(doc, DATA["footer"], size=9, color=MUTED, space_before=16)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / f"Afrivate-Core-Team-Engagement-Equity-Letter-{letter['slug']}.docx"
    dl = Path(rf"C:\Users\DELL\Downloads\Afrivate Engagement Equity Letter - {letter['fullName']}.docx")
    doc.save(out)
    copyfile(out, dl)
    print("saved", out)
    print("copied", dl)


for letter in DATA["letters"]:
    write_letter(letter)
