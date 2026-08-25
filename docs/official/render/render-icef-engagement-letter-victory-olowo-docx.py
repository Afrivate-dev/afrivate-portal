"""Branded DOCX companion for the Victory Olowo ICEF engagement letter."""
from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

PURPLE = RGBColor(0x8D, 0x40, 0x87)
INK = RGBColor(0x1F, 0x1F, 0x1F)
MUTED = RGBColor(0x5F, 0x5F, 0x5F)
SOFT = "F8F3F8"

ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "brand" / "afrivate-logo-long-purple.png"
OUT = ROOT / "hiring" / "engagement-letters" / "Afrivate-ICEF-Engagement-Letter-Victory-Olowo.docx"
DL = Path(r"C:\Users\DELL\Downloads\Afrivate ICEF Engagement Letter - Victory Olowo.docx")


def set_run(run, *, size=11, bold=False, color=INK, all_caps=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.font.all_caps = all_caps


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def no_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def add_bottom_border(paragraph, color, size):
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    paragraph._p.get_or_add_pPr().append(pBdr)


def para(doc, text, *, size=11, bold=False, color=INK, space_after=8, space_before=0, align="left", all_caps=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, all_caps=all_caps)
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    add_bottom_border(p, "EBDCEB", "6")
    run = p.add_run(text)
    set_run(run, size=11, bold=True, color=PURPLE, all_caps=True)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run(run, size=11)
    return p


doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    r = fp.add_run("hr@afrivate.org  ·  portal.afrivate.org          AfriVate Technologies Ltd  ·  RC: 9210092")
    set_run(r, size=8, color=MUTED)

header = doc.add_table(rows=1, cols=2)
no_table_borders(header)
c0, c1 = header.rows[0].cells
c0.paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(1.7))
c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
for i, line in enumerate(["Official Document", "AfriVate Technologies Ltd", "RC: 9210092"]):
    target = c1.paragraphs[0] if i == 0 else c1.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    target.paragraph_format.space_after = Pt(0)
    target.paragraph_format.space_before = Pt(0)
    set_run(target.add_run(line), size=9, color=MUTED)

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(4)
rule.paragraph_format.space_after = Pt(10)
add_bottom_border(rule, "8D4087", "18")

para(doc, "Internal Contributor Engagement Letter", size=16, bold=True, align="center", space_after=2, all_caps=True)
para(
    doc,
    "Unpaid Placement — Virtual Assistant, Human Resources (People & Culture)",
    size=10,
    color=MUTED,
    align="center",
    space_after=12,
)

meta = [
    ("To", "Victory Olowo"),
    ("From", "Joshua Oluwasujibomi Komolafe, Chief Executive Officer, AfriVate Technologies Ltd"),
    ("Role", "Virtual Assistant — Human Resources (Internal Contributor)"),
    ("Reports to", "Emmanuel Okpiaifo, Human Resources Manager (People & Culture)"),
    ("Location", "Remote"),
    ("Start Date", "5 August 2026"),
    ("Date", "20 August 2026"),
    ("Document Reference", "Issued under AFRI-ICEF-01 (Internal Contributor Engagement Framework)"),
]
mt = doc.add_table(rows=len(meta), cols=2)
no_table_borders(mt)
for i, (k, v) in enumerate(meta):
    a, b = mt.rows[i].cells
    shade_cell(a, SOFT)
    shade_cell(b, SOFT)
    pa = a.paragraphs[0]
    pa.paragraph_format.space_after = Pt(2)
    pa.paragraph_format.space_before = Pt(2)
    set_run(pa.add_run(k), size=10, bold=True)
    pb = b.paragraphs[0]
    pb.paragraph_format.space_after = Pt(2)
    pb.paragraph_format.space_before = Pt(2)
    set_run(pb.add_run(v), size=10, color=MUTED)

doc.add_paragraph()
para(doc, "Dear Victory,", size=11, bold=True, space_after=8)
para(
    doc,
    "AfriVate Technologies Ltd (“AfriVate”) is pleased to confirm your placement as a Virtual Assistant supporting the Human Resources function within the People & Culture pillar. This letter confirms your engagement as an Internal Contributor and records the specific terms applicable to your placement. This letter is a written instrument for the purposes of AFRI-ICEF-01; to the extent of any direct conflict between this letter and AFRI-ICEF-01 on a matter this letter expressly addresses, this letter governs. On all other matters, and wherever this letter is silent, AFRI-ICEF-01 and the other policies listed in Clause 10 apply. Please read this letter together with those documents, all of which are available to you in the AfriVate Portal under Resources.",
)

heading(doc, "1. Nature of This Engagement")
para(
    doc,
    "This is an unpaid internal contribution and is not a contract of employment. In line with AFRI-ICEF-01, nothing in this letter, in any title used, in Portal access, or in the performance of work creates employment, a contract of service, partnership, joint venture, or agency, or any entitlement to salary, wages, benefits, or equity, unless a separate written instrument signed by the CEO expressly creates that entitlement.",
)
para(
    doc,
    "Your engagement is governed exclusively by AFRI-ICEF-01, and this letter records the terms specific to your placement within that Framework. Where this letter is silent, AFRI-ICEF-01 and the other policies referenced in Clause 10 apply.",
)

heading(doc, "2. Duration")
para(
    doc,
    "Your placement commences on the Start Date stated above. This is an indefinite, ongoing placement with no fixed end date — it is not an internship with a set end date, and no minimum or maximum duration is promised by either party. Either you or AfriVate may end it at any time in accordance with Clause 9 below. Your placement will be reviewed by your Team Lead and People & Culture no less than every six (6) months, consistent with the review cycle set out in AFRI-ICEF-01, to confirm your Agreed Capacity, responsibilities, and standing remain appropriate; this review is a checkpoint, not a renewal requirement, and your placement continues automatically unless it is ended under Clause 9. AfriVate may transition you to a paid role at its sole discretion under Clause 9 (Path to Paid Employment) of AFRI-ICEF-01, or agree with you in writing to change your responsibilities or Agreed Capacity at any time.",
)

heading(doc, "3. Working Days and Agreed Capacity")
para(
    doc,
    "Your official work days are Monday to Thursday, consistent with the official work week that applies across AfriVate. AfriVate does not operate fixed working hours; consistent with the definition of Agreed Capacity in AFRI-ICEF-01, your Agreed Capacity will be expressed as deliverables and outcomes rather than a fixed hour count. Your Team Lead, Emmanuel Okpiaifo, will agree specific tasks with you, each recorded in Portal → Tasks with a clear outcome, priority, and deadline, per AFRI-SWP; you are expected to meet agreed deadlines rather than log a set number of hours. During official work days, you must still be reachable on Slack during the core hours agreed with your Team Lead, and must acknowledge official Slack messages within four (4) hours during those hours, per AFRI-SWP. Deadlines will be set at a level reasonable for an unpaid placement of this nature, undertaken alongside your other commitments, and will take account of those commitments; you should raise any concern about workload or a deadline with your Team Lead as soon as it arises, and doing so will not affect your standing on the placement. AfriVate does not direct the specific hours or location in which you complete your tasks, and nothing in this Clause requires you to be available outside the core hours agreed for official work days.",
)

heading(doc, "4. Key Responsibilities")
for item in [
    "Support the Human Resources Manager with daily administrative tasks, including scheduling, correspondence, and calendar coordination.",
    "Assist with onboarding coordination for new Team Members and Internal Contributors, including tracking Portal onboarding checklist completion.",
    "Help organize and maintain HR files, trackers, and Portal documentation, including candidate and contributor records.",
    "Draft and proofread routine HR communications, job posts, and memos under supervision, for review and approval by the HR Manager before use.",
    "Support scheduling and logistics for interviews, onboarding sessions, and other HR meetings.",
    "Assist in tracking policy acknowledgement completion in Portal → Resources and following up with Team Members as directed.",
    "Support administration of engagement surveys, learning assignment tracking, and other Portal → People workflows as directed.",
    "Take notes and log action items for HR meetings, and help maintain accurate records of HR decisions in the Portal.",
    "Provide general virtual assistant support to the Human Resources Manager as reasonably required.",
]:
    bullet(doc, item)
para(
    doc,
    "All tasks are assigned and tracked in Portal → Tasks per AFRI-SWP. As an Internal Contributor at Decision Class D under AFRI-DOA-01, your role is to execute assigned work accurately — you do not have authority to approve leave, warnings, hiring decisions, or any Cash Commitment; these remain with People & Culture, your Team Lead, or the CEO as applicable. Responsibilities may evolve as the placement progresses.",
)

heading(doc, "5. Compensation and Support")
para(
    doc,
    "This is an unpaid placement and does not constitute employment; no element of this Clause creates a salary, wage, or fixed entitlement. As a specific exception approved in writing by the CEO's signature on this letter, under AFRI-DOA-01 (Cash Commitment approval) and AFRI-ICEF-01 Clause 4, AfriVate approves a discretionary data support reimbursement of up to ₦5,000 (five thousand naira) per month toward your connectivity costs for the duration of this placement. This is not a standing or automatic payment: each month, you may submit a request with reasonable supporting information to your Team Lead, who will forward it to Finance & Administration for review and logging per AFRI-DOA-01; Finance may approve, reduce, or decline any individual request, including where AfriVate's funding position changes. AfriVate may vary or withdraw this discretionary reimbursement at any time without this constituting a breach of this letter, and its provision in any given month creates no entitlement to it in any other month.",
)
para(doc, "In addition to the discretionary data support reimbursement, your placement includes:")
for item in [
    "Structured supervision and regular feedback from the Human Resources Manager.",
    "Priority access to learning assignments and courses made available through Portal → Learning.",
    "Direct mentorship from the Human Resources Manager on HR processes and professional development.",
    "A letter of recommendation on satisfactory completion of the placement, per Clause 6 below.",
    "A factual reference (role, dates, and nature of contribution) available on request, per AFRI-ICEF-01.",
    "Consideration for future paid roles at AfriVate, if and when capacity opens, per Clause 11 below.",
]:
    bullet(doc, item)
para(
    doc,
    "No other pay, equity, or benefit is promised beyond what is stated in this letter or in a separate written instrument signed by the CEO, per AFRI-ICEF-01 Clause 10.",
)

heading(doc, "6. Letter of Recommendation")
para(
    doc,
    "AFRI-ICEF-01 entitles every Internal Contributor in good standing, on request and subject to proper handover, to a factual reference limited to role, dates, and nature of contribution. In addition to that baseline entitlement, and as a specific written commitment made binding by the CEO's signature on this letter, AfriVate will issue you a signed letter of recommendation addressing your role, responsibilities, tenure, and performance, at the earlier of: (a) your written request, provided you have then completed at least three (3) months of active placement in good standing; or (b) the end of your placement under Clause 9, provided your engagement is not ended by AfriVate for serious breach under Clause 9(b) and you complete Portal handover as required by AFRI-SWP. Your Team Lead will make this determination based on your Portal task record, weekly check ins, and any documented coaching or warnings on file. If you disagree with the determination, you may ask People & Culture to review it, and People & Culture's decision on review is final. This additional commitment applies only to this placement and does not itself create any right to pay, equity, or continued engagement.",
)

heading(doc, "7. Confidentiality and Data Protection")
para(
    doc,
    "You will have access to Confidential Information as defined in AFRI-ICEF-01, including personal data of employees, candidates, and Internal Contributors handled through HR. You must protect this information indefinitely, during and after your placement, and must not disclose it outside AfriVate without prior written authorisation. You must handle all personal data you process in this role in accordance with the Nigeria Data Protection Act 2023 — accessing only what your assigned tasks require, keeping it confidential, and using only AfriVate's approved systems to store or share it. Nothing in this Clause prevents you from: reporting suspected misconduct in good faith through Portal → People → Speak up or to a relevant regulator or law enforcement authority; making a disclosure required by law or by a court, tribunal, or regulator of competent jurisdiction; or seeking confidential legal advice about your rights under this letter.",
)

heading(doc, "8. Intellectual Property")
para(
    doc,
    "In consideration of being offered this placement and the benefits described in Clause 5, you agree that all Work Product you create in the course of this placement — including documents, templates, and processes — is owned exclusively by AfriVate from creation, and you assign to AfriVate all right, title, and interest in it worldwide, per AFRI-ICEF-01 Clause 6. This does not apply to work you already owned before this placement began, or work created entirely outside your duties and without AfriVate's time, resources, or confidential information.",
)

heading(doc, "9. Ending the Engagement")
bullet(
    doc,
    "By you: at least fourteen (14) calendar days' written notice to your Team Lead and hr@afrivate.org, with Portal handover completed, per AFRI-ICEF-01 Clause 8.",
)
bullet(
    doc,
    "By AfriVate: immediately, without notice, for serious breach (including confidentiality, IP, dishonesty, harassment, or security violations); otherwise, ordinarily seven (7) calendar days' written notice as a courtesy standard, not a right to continued access.",
)
bullet(
    doc,
    "On ending the placement, however it ends: your Portal and Slack access will be revoked, and your confidentiality and IP obligations continue to apply.",
)

heading(doc, "10. Governing Policies")
para(
    doc,
    "This placement is governed by AFRI-ICEF-01 together with: AFRI-SWP (Standard Work Process), AFRI-ORG-01 (Organisational Structure), AFRI-LAP-01 (Leave and Absence Policy), AFRI-DOA-01 (Delegation of Authority), and AFRI-EOH-01 (Team Member Onboarding Handbook). You must acknowledge each applicable document in Portal → Resources within seven (7) official work days of your access being approved. Any absence must be requested through the Portal under AFRI-LAP-01, with at least three (3) official work days' notice for ordinary planned absence.",
)

heading(doc, "11. Not a Promise of Future Employment")
para(
    doc,
    "This placement does not guarantee you a paid position at AfriVate afterward. Prior unpaid contribution is a factor AfriVate may consider if a paid role becomes available, but it creates no automatic right to an offer, salary level, or equity, per AFRI-ICEF-01 Clause 9.",
)

heading(doc, "12. Effect if This Relationship Is Later Recharacterised")
para(
    doc,
    "Both parties intend this placement to be an unpaid internal contribution under AFRI-ICEF-01, and not a contract of employment, and have structured this letter on that basis in good faith. If a court, tribunal, or statutory authority of competent jurisdiction nonetheless determines that a different relationship existed in law for a given period, this letter is deemed varied to the minimum extent required to comply with the applicable mandatory legal requirements for that period, and the remainder of this letter continues in force. Such a determination does not, by itself, entitle you to any benefit, payment, or status beyond what that determination legally requires, and does not affect any period before or after the period to which the determination relates.",
)

heading(doc, "13. Acceptance")
para(
    doc,
    "Please sign below to confirm that you have read, understood, and accept the terms of this Internal Contributor placement, together with AFRI-ICEF-01 and the other policies referenced above.",
)
para(doc, "Sincerely,", space_after=18)

sig = doc.add_table(rows=1, cols=2)
no_table_borders(sig)
left, right = sig.rows[0].cells
set_run(left.paragraphs[0].add_run("Joshua Oluwasujibomi Komolafe"), size=11, bold=True)
for line in ["Chief Executive Officer", "AfriVate Technologies Ltd"]:
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(line), size=10, color=MUTED)

set_run(right.paragraphs[0].add_run("Accepted by: Victory Olowo"), size=11, bold=True)
for label in ["Signature", "Date"]:
    p = right.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("________________________________"), size=11, color=MUTED)
    p2 = right.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_run(p2.add_run(label), size=9, color=MUTED)

para(
    doc,
    "Document Reference: Issued under AFRI-ICEF-01 · Related: AFRI-SWP · AFRI-ORG-01 · AFRI-LAP-01 · AFRI-DOA-01 · AFRI-EOH-01 · Governed by the laws of the Federal Republic of Nigeria; courts of the Federal Capital Territory, Abuja have jurisdiction, subject to any mandatory forum rules.",
    size=9,
    color=MUTED,
    space_before=16,
)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
copyfile(OUT, DL)
print("saved", OUT)
print("copied", DL)
